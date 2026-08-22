import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import io
from datetime import datetime

# DOCX व PDF जनरेशनसाठी आवश्यक लायब्ररीज
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("कृपया 'python-docx' लायब्ररी इन्स्टॉल करा: pip install python-docx")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfgen import canvas
except ImportError:
    st.error("कृपया 'reportlab' लायब्ररी इन्स्टॉल करा: pip install reportlab")

# ==============================================================================
# १. ॲप्लिकेशन कॉन्फिगरेशन व सेशन स्टेट (Configuration & Session State)
# ==============================================================================
st.set_page_config(
    page_title="आकांक्षा AI कायदेशीर महा-सहाय्यक",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# सेशन स्टेट इनिशियालायझेशन
if "current_page" not in st.session_state:
    st.session_state.current_page = "गृहपृष्ठ (Home)"
if "draft_history" not in st.session_state:
    st.session_state.draft_history = []
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "font_size" not in st.session_state:
    st.session_state.font_size = 14
if "page_margin" not in st.session_state:
    st.session_state.page_margin = 15

# ==============================================================================
# २. कस्टम CSS, थीम व ब्रँडिंग (Custom CSS Styling)
# ==============================================================================
st.markdown("""
<style>
    /* Google Fonts - Mukta */
    @import url('https://fonts.googleapis.com/css2?family=Mukta:wght@300;400;600;700;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Mukta', sans-serif;
        background-color: #0F172A;
        color: #E2E8F0;
    }
    
    /* मुख्य शीर्षक - ॲनिमेटेड ग्लॉसी टेक्स्ट */
    .main-title {
        font-size: 2.6rem;
        font-weight: 800;
        text-align: center;
        background: linear-gradient(90deg, #FF4B4B, #FF8C00, #00E676, #00B0FF, #D500F9);
        background-size: 300% 300%;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: glowGradient 6s ease infinite;
        margin-bottom: 5px;
    }
    
    @keyframes glowGradient {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    
    .sub-title {
        text-align: center;
        color: #94A3B8;
        font-size: 1.2rem;
        margin-bottom: 20px;
        font-weight: 600;
    }
    
    /* VIP बॅनर डिझाइन */
    .vip-banner {
        background: rgba(15, 23, 42, 0.8);
        border: 2px solid transparent;
        border-radius: 12px;
        padding: 12px;
        text-align: center;
        font-size: 1.1rem;
        font-weight: 700;
        color: #FFFFFF;
        box-shadow: 0 0 15px rgba(0, 230, 118, 0.4);
        animation: vipPulse 3s infinite alternate;
        margin-bottom: 25px;
    }
    
    @keyframes vipPulse {
        0% {
            border-color: #00E676;
            box-shadow: 0 0 10px #00E676, inset 0 0 10px #00E676;
        }
        50% {
            border-color: #00B0FF;
            box-shadow: 0 0 15px #00B0FF, inset 0 0 15px #00B0FF;
        }
        100% {
            border-color: #FF4B4B;
            box-shadow: 0 0 20px #FF4B4B, inset 0 0 20px #FF4B4B;
        }
    }
    
    /* कार्ड्स डिझाइन */
    .card-button {
        background: linear-gradient(135deg, #1E293B 0%, #334155 100%);
        border: 1px solid #475569;
        border-radius: 15px;
        padding: 20px;
        text-align: center;
        transition: all 0.3s ease;
        cursor: pointer;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin-bottom: 15px;
    }
    .card-button:hover {
        transform: translateY(-5px);
        border-color: #38BDF8;
        box-shadow: 0 10px 25px -5px rgba(56, 189, 248, 0.4);
    }
    .card-icon {
        font-size: 2.5rem;
        margin-bottom: 10px;
    }
    .card-title {
        font-size: 1.2rem;
        font-weight: 700;
        color: #F8FAFC;
    }
    
    /* चॅट बबल्स */
    .user-bubble {
        background-color: #0284C7;
        color: white;
        padding: 12px 16px;
        border-radius: 18px 18px 2px 18px;
        margin: 8px 0;
        max-width: 80%;
        float: right;
        clear: both;
        font-size: 1.05rem;
    }
    .ai-bubble {
        background-color: #334155;
        color: #F1F5F9;
        padding: 12px 16px;
        border-radius: 18px 18px 18px 2px;
        margin: 8px 0;
        max-width: 85%;
        float: left;
        clear: both;
        border-left: 4px solid #38BDF8;
        font-size: 1.05rem;
    }
    
    /* मसुदा बॉक्स (Draft Output Box) */
    .draft-container {
        background-color: #FFFFFF;
        color: #000000;
        padding: 30px;
        border-radius: 8px;
        box-shadow: 0 5px 15px rgba(0,0,0,0.3);
        font-family: 'Mukta', sans-serif;
        line-height: 1.6;
        margin-top: 20px;
        margin-bottom: 20px;
        white-space: pre-wrap;
    }
    
    /* कस्टम बटणे */
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        font-weight: 700;
        background: linear-gradient(90deg, #2563EB, #1D4ED8);
        color: white;
        border: none;
        padding: 10px 16px;
        transition: all 0.2s;
    }
    .stButton>button:hover {
        background: linear-gradient(90deg, #1D4ED8, #1E40AF);
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.4);
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# ३. हेल्पर फंक्शन्स (PDF, Word, Email, WhatsApp, Draft Logic)
# ==============================================================================

def generate_docx_file(text_content, title="कायदेशीर_मसुदा"):
    """मसुद्याची Word (.docx) फाईल तयार करते"""
    doc = docx.Document()
    
    # पेज मार्जिन सेटिंग्स
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    p = doc.add_paragraph()
    run = p.add_run(text_content)
    run.font.name = 'Arial'
    run.font.size = Pt(st.session_state.font_size)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_file(text_content, title="कायदेशीर_मसुदा"):
    """मसुद्याची PDF (.pdf) फाईल तयार करते"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=st.session_state.page_margin,
        leftMargin=st.session_state.page_margin,
        topMargin=st.session_state.page_margin,
        bottomMargin=st.session_state.page_margin
    )
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        'MarathiStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=st.session_state.font_size,
        leading=st.session_state.font_size + 6
    )
    
    story = []
    paragraphs = text_content.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para, normal_style))
        else:
            story.append(Spacer(1, 10))
            
    doc.build(story)
    buffer.seek(0)
    return buffer

def get_whatsapp_share_link(text):
    """WhatsApp शेअर लिंक तयार करणे"""
    encoded_text = urllib.parse.quote(text[:1500] + "\n\n(संपूर्ण मसुदा आकांक्षा AI द्वारे तयार केला आहे)")
    return f"https://api.whatsapp.com/send?text={encoded_text}"

def get_email_share_link(subject, body, to_email=""):
    """E-mail शेअर लिंक तयार करणे"""
    encoded_subject = urllib.parse.quote(subject)
    encoded_body = urllib.parse.quote(body)
    return f"mailto:{to_email}?subject={encoded_subject}&body={encoded_body}"

def save_to_draft_history(draft_type, content):
    """मसुदा इतिहासात सेव्ह करणे"""
    timestamp = datetime.datetime.now().strftime("%d-%m-%Y %H:%M")
    st.session_state.draft_history.append({
        "id": len(st.session_state.draft_history) + 1,
        "type": draft_type,
        "date": timestamp,
        "content": content
    })
    st.toast(f"✅ {draft_type} मसुदा यशस्वीरीत्या जतन (Save) झाला!", icon="💾")

def render_draft_action_buttons(draft_type, text_content, email_target=""):
    """मसुदा तयार झाल्यावर डाऊनलोड व शेअरिंग बटणे दाखवणे"""
    st.markdown("---")
    st.subheader("📥 मसुदा डाऊनलोड व शेअरिंग ऑप्शन्स")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        docx_file = generate_docx_file(text_content, title=draft_type)
        st.download_button(
            label="📄 Word (.docx)",
            data=docx_file,
            file_name=f"{draft_type}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    with col2:
        pdf_file = generate_pdf_file(text_content, title=draft_type)
        st.download_button(
            label="🔴 PDF (.pdf)",
            data=pdf_file,
            file_name=f"{draft_type}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
            mime="application/pdf"
        )
        
    with col3:
        wa_link = get_whatsapp_share_link(text_content)
        st.markdown(f'<a href="{wa_link}" target="_blank"><button style="width:100%; background-color:#25D366; color:white; border:none; padding:10px; border-radius:8px; font-weight:bold;">💬 WhatsApp शेअर</button></a>', unsafe_allow_html=True)
        
    with col4:
        email_link = get_email_share_link(f"कायदेशीर मसुदा - {draft_type}", text_content, email_target)
        st.markdown(f'<a href="{email_link}" target="_blank"><button style="width:100%; background-color:#EA4335; color:white; border:none; padding:10px; border-radius:8px; font-weight:bold;">✉️ Email द्वारे पाठवा</button></a>', unsafe_allow_html=True)
        
    st.write("")
    if st.button("💾 हा मसुदा इतिहासात (History) जतन करा", key=f"save_{draft_type}_{datetime.datetime.now().microsecond}"):
        save_to_draft_history(draft_type, text_content)

# ==============================================================================
# ४. हेडर्स व साईडबार (Headers & Sidebar Navigation)
# ==============================================================================

# ॲप मुख्य शीर्षक
st.markdown('<div class="main-title">आकांक्षा AI कायदेशीर व प्रशासकीय महा-सहाय्यक</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">⚡ एका सेकंदात अर्ज व कायदेशीर मसुदे तयार करा!</div>', unsafe_allow_html=True)

# VIP बॅनर
st.markdown('''
<div class="vip-banner">
    👨‍💼 संकल्पना व निर्मिती: <b>सतीश अशोक प्रधान</b> | 📱 मो. <b>८६६८२३५३९५</b>
</div>
''', unsafe_allow_html=True)

# साईडबार नेव्हिगेशन
st.sidebar.title("📌 मुख्य नेव्हिगेशन")

page_options = [
    "गृहपृष्ठ (Home)",
    "१) जोडपत्र 'अ' (RTI 6(1))",
    "२) जोडपत्र 'ब' (प्रथम अपील 19(1))",
    "३) जोडपत्र 'क' (द्वितीय अपील 19(3))",
    "४) ✨ AI चॅट (आकांक्षा AI)",
    "५) 📜 न्यायालयीन मसुदा",
    "६) 📢 शासकीय तक्रार अर्ज",
    "७) 📝 प्रतिज्ञापत्र (Affidavit)",
    "८) 🛒 ग्राहक मंच तक्रार",
    "📚 कायदेशीर संदर्भ ग्रंथालय",
    "💾 माझे जतन केलेले मसुदे"
]

selected_sidebar = st.sidebar.radio(
    "विभाग निवडा:",
    page_options,
    index=page_options.index(st.session_state.current_page) if st.session_state.current_page in page_options else 0
)

if selected_sidebar != st.session_state.current_page:
    st.session_state.current_page = selected_sidebar
    st.rerun()

# साईडबार मधील प्रिंट व फॉन्ट सेटिंग्ज
st.sidebar.markdown("---")
st.sidebar.subheader("🖨️ प्रिंट व फॉन्ट सेटिंग्ज")
st.session_state.font_size = st.sidebar.slider("फॉन्ट साईझ (Pt):", 10, 24, 14)
st.session_state.page_margin = st.sidebar.slider("पेज मार्जिन (mm):", 5, 30, 15)

st.sidebar.markdown("---")
st.sidebar.info("💡 **टीप:** तयार झालेला मसुदा थेट A4 साईझ मध्ये प्रिंट किंवा PDF/Word द्वारे डाउनलोड करता येतो.")

# ==============================================================================
# ५. विभाग १: गृहपृष्ठ (Home - 8 Grid Cards)
# ==============================================================================

if st.session_state.current_page == "गृहपृष्ठ (Home)":
    st.subheader("🎯 त्वरित सेवेसाठी खालीलपैकी एका पर्यायावर क्लिक करा:")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 जोडपत्र 'अ' (माहिती अधिकार मूळ अर्ज - कलम ६(१))", key="card1"):
            st.session_state.current_page = "१) जोडपत्र 'अ' (RTI 6(1))"
            st.rerun()
            
        if st.button("⚖️ जोडपत्र 'ब' (प्रथम अपील - कलम १९(१))", key="card2"):
            st.session_state.current_page = "२) जोडपत्र 'ब' (प्रथम अपील 19(1))"
            st.rerun()
            
        if st.button("🏛️ जोडपत्र 'क' (द्वितीय अपील - राज्य माहिती आयोग)", key="card3"):
            st.session_state.current_page = "३) जोडपत्र 'क' (द्वितीय अपील 19(3))"
            st.rerun()
            
        if st.button("✨ AI चॅट (आकांक्षा AI कायदेशीर सहाय्यक)", key="card4"):
            st.session_state.current_page = "४) ✨ AI चॅट (आकांक्षा AI)"
            st.rerun()

    with col2:
        if st.button("📜 न्यायालयीन मसुदा (दिवाणी/फौजदारी/रिट याचिका)", key="card5"):
            st.session_state.current_page = "५) 📜 न्यायालयीन मसुदा"
            st.rerun()
            
        if st.button("📢 शासकीय तक्रार अर्ज (प्रशासकीय गैरकारभाराविरुद्ध)", key="card6"):
            st.session_state.current_page = "६) शासकीय तक्रार अर्ज"
            st.rerun()
            
        if st.button("📝 प्रतिज्ञापत्र (Affidavit Draft)", key="card7"):
            st.session_state.current_page = "७) 📝 प्रतिज्ञापत्र (Affidavit)"
            st.rerun()
            
        if st.button("🛒 ग्राहक मंच तक्रार (Consumer Forum)", key="card8"):
            st.session_state.current_page = "८) 🛒 ग्राहक मंच तक्रार"
            st.rerun()

    st.markdown("---")
    st.markdown("""
    ### 🌟 आकांक्षा AI ची मुख्य वैशिष्ट्ये:
    * 📑 **परिफूर्ण कायदेशीर स्वरूप:** महाराष्ट्र माहिती अधिकार नियम व भारतीय कायद्यांनुसार कायदेशीर मसुदा.
    * ⚡ **झटपट जनरेशन:** सर्व अर्जांचे फॉर्म सोपे व जलद भरण्यासाठी डिझाइन केलेले आहेत.
    * 📤 **थेट शेअरिंग:** WhatsApp, E-mail, PDF आणि Word डाउनलोडची सोय.
    * 🔒 **पूर्ण गोपनीयता:** तुमचा डेटा तुमच्या स्थानिक सेशन्समध्ये सुरक्षित राहतो.
    """)

# ==============================================================================
# ६. विभाग २: जोडपत्र 'अ' (RTI 6(1) Original Application)
# ==============================================================================

elif st.session_state.current_page == "१) जोडपत्र 'अ' (RTI 6(1))":
    st.header("📄 जोडपत्र 'अ' - माहिती अधिकार नियम २००५ कलम ६(१) अन्वये अर्ज")
    
    with st.form("rti_a_form"):
        col1, col2 = st.columns(2)
        with col1:
            app_name = st.text_input("अर्जादाराचे नाव:", value="सतीश अशोक प्रधान")
            app_address = st.text_area("अर्जादाराचा पूर्ण पत्ता:", value="छत्रपती संभाजीनगर")
            app_mobile = st.text_input("मोबाईल क्रमांक:", value="८६६८२३५३९५")
        with col2:
            pio_office = st.text_area("जन माहिती अधिकाऱ्याचे कार्यालय / विभाग / पत्ता:", placeholder="उदा. जन माहिती अधिकारी, जिल्हाधिकारी कार्यालय, छत्रपती संभाजीनगर")
            subject = st.text_input("माहितीचा विषय:", placeholder="उदा. रस्ता कामाचा निधी व निविदा बाबत...")
            period = st.text_input("माहितीचा कालावधी (उदा. २०२३ ते २०२५):", placeholder="उदा. १ जानेवारी २०२३ ते ३१ डिसेंबर २०२४")
        
        info_details = st.text_area("हवी असलेल्या माहितीचे मुद्देसूद वर्णन (प्रत्येक मुद्दा नवीन ओळीवर):", placeholder="१. निविदेची प्रत\n२. खर्च झालेला निधी\n३. मोजमाप पुस्तकाची (MB Book) प्रमाणित प्रत")
        
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            delivery_mode = st.radio("माहिती कशी हवी आहे?", ["व्यक्तिशः (Self)", "टपालाद्वारे (नोंदणीकृत / स्पीड पोस्ट)", "जी-मेल (Gmail/E-mail ID वर)"])
            user_email = ""
            if "जी-मेल" in delivery_mode:
                user_email = st.text_input("ई-मेल आयडी प्रविष्ट करा:", placeholder="example@gmail.com")
        
        with col_m2:
            bpl_status = st.radio("अर्जादार दारिद्र्यरेषेखालील (BPL) आहे का?", ["नाही (Rs. 10 कोर्ट फी स्टॅम्प आवश्यक)", "होय (BPL कार्डधारक - फी माफ)"])
            bpl_card_no = ""
            if "होय" in bpl_status:
                bpl_card_no = st.text_input("BPL कार्ड / रेशन कार्ड क्रमांक:")

        submit_a = st.form_submit_button("🚀 जोडपत्र 'अ' मसुदा तयार करा")

    if submit_a:
        delivery_str = delivery_mode
        if "जी-मेल" in delivery_mode:
            delivery_str += f" ({user_email})"

        bpl_str = "नाही. (रू. १०/- चा कोर्ट फी स्टॅम्प / चलनाची प्रत जोडली आहे)"
        if "होय" in bpl_status:
            bpl_str = f"होय. (माझे दारिद्र्यरेषेखालील कार्ड क्र. {bpl_card_no} ची प्रत जोडली आहे, त्यामुळे फी लागू होत नाही)"

        rti_a_text = f"""परिशिष्ट / जोडपत्र 'अ'
(नियम ३ पहा)
माहितीचा अधिकार अधिनियम, २००५ याच्या कलम ६ (१) खालील माहिती मिळण्यासाठीचा अर्ज.

प्रति,
श्री. जन माहिती अधिकारी,
{pio_office}

१. अर्जदाराचे पूर्ण नाव      : {app_name}
२. अर्जदाराचा पत्ता        : {app_address}
   मोबाईल क्र.           : {app_mobile}

३. हवी असलेल्या माहितीचा तपशील:
   (क) माहितीचा विषय     : {subject}
   (ख) माहितीचा कालावधी  : {period}
   (ग) हव्या असलेल्या माहितीचे वर्णन:
{info_details}

४. माहिती टपालाने हवी आहे की व्यक्तिशः: {delivery_str}

५. अर्जदार दारिद्र्यरेषेखालील आहे किंवा कसे: {bpl_str}

६. अर्जाचे शुल्क:
   माहिती अधिकार (नियम) नुसार आवश्यक ते रु. १०/- चे न्यायालयीन फी तिकीट (Court Fee Stamp) अर्जावर लावले आहे / ऑनलाईन भरले आहे.

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {app_name} )
                                                     अर्जादाराची स्वाक्षरी"""

        st.success("✅ जोडपत्र 'अ' मसुदा यशस्वीरीत्या तयार झाला आहे!")
        st.markdown(f'<div class="draft-container">{rti_a_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("जोडपत्र_अ_RTI", rti_a_text)

# ==============================================================================
# ७. विभाग ३: जोडपत्र 'ब' (First Appeal 19(1))
# ==============================================================================

elif st.session_state.current_page == "२) जोडपत्र 'ब' (प्रथम अपील 19(1))":
    st.header("⚖️ जोडपत्र 'ब' - प्रथम अपील अर्ज (कलम १९(१))")
    
    with st.form("rti_b_form"):
        col1, col2 = st.columns(2)
        with col1:
            app_name = st.text_input("अपीलकर्त्याचे नाव:", value="सतीश अशोक प्रधान")
            app_address = st.text_area("अपीलकर्त्याचा पत्ता:", value="छत्रपती संभाजीनगर")
            app_mobile = st.text_input("मोबाईल क्रमांक:", value="८६६८२३५३९५")
        with col2:
            first_aa = st.text_area("प्रथम अपीलीय अधिकाऱ्याचे पद व पत्ता:", placeholder="उदा. मा. प्रथम अपीलीय अधिकारी तथा अपर जिल्हाधिकारी, छत्रपती संभाजीनगर")
            pio_details = st.text_area("जन माहिती अधिकाऱ्याचे नाव व कार्यालय:", placeholder="उदा. जन माहिती अधिकारी, तहसीलदार कार्यालय...")
        
        orig_date = st.text_input("मूळ अर्ज (जोडपत्र 'अ') सादर केल्याची तारीख:", placeholder="उदा. १५/०१/२०२५")
        reason = st.selectbox("प्रथम अपीलाचे मुख्य कारण:", [
            "मुदतीत (३० दिवसांत) कोणतीही माहिती मिळाली नाही.",
            "जन माहिती अधिकाऱ्याने दिलेली माहिती अपूर्ण व दिशाभूल करणारी आहे.",
            "माहिती नाकारण्यात आली आहे.",
            "अवाजवी शुल्काची मागणी करण्यात आली आहे."
        ])
        prayer = st.text_area("मागणी (Prayer):", value="माहिती अधिकार कायदा कलम ७(६) नुसार मुदतीत माहिती न दिल्यामुळे संपूर्ण माहिती विनामूल्य तात्काळ पुरवण्यात यावी व दोषी अधिकाऱ्यावर कारवाई व्हावी.")
        
        submit_b = st.form_submit_button("🚀 जोडपत्र 'ब' प्रथम अपील तयार करा")

    if submit_b:
        rti_b_text = f"""परिशिष्ट / जोडपत्र 'ब'
[ नियम ५ (१) पहा ]
माहितीचा अधिकार अधिनियम, २००५ च्या कलम १९(१) अन्वये प्रथम अपीलाचा अर्ज.

प्रति,
मा. प्रथम अपीलीय अधिकारी,
{first_aa}

१. अपीलकर्त्याचे नाव        : {app_name}
२. अपीलकर्त्याचा पत्ता      : {app_address}
   मोबाईल क्र.             : {app_mobile}

३. जन माहिती अधिकाऱ्याचा तपशील: {pio_details}

४. मूळ अर्ज जोडपत्र 'अ' सादर केल्याचा दिनांक: {orig_date}

५. अपीलाचे कारण           : {reason}

६. अपीलाचा सविस्तर तपशील व वस्तुस्थिती:
   मी वरील नमूद तारखेस जन माहिती अधिकाऱ्यांकडे रीतसर अर्ज सादर केला होता. परंतु, {reason} यामुळे मी हे प्रथम अपील दाखल करत आहे.

७. अपेक्षित आदेश / मागणी (Prayer):
   {prayer}

८. अर्जासोबत जोडलेली कागदपत्रे:
   १. मूळ अर्ज (जोडपत्र 'अ') ची प्रत.
   २. टपाल पोहोच / पोच पावती.
   ३. जन माहिती अधिकाऱ्याचे पत्र (असल्यास).

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {app_name} )
                                                     अपीलकर्त्याची स्वाक्षरी"""

        st.success("✅ जोडपत्र 'ब' प्रथम अपील तयार झाले आहे!")
        st.markdown(f'<div class="draft-container">{rti_b_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("जोडपत्र_ब_प्रथम_अपील", rti_b_text)

# ==============================================================================
# ८. विभाग ४: जोडपत्र 'क' (Second Appeal 19(3))
# ==============================================================================

elif st.session_state.current_page == "३) जोडपत्र 'क' (द्वितीय अपील 19(3))":
    st.header("🏛️ जोडपत्र 'क' - द्वितीय अपील (राज्य माहिती आयोग कलम १९(३))")
    
    with st.form("rti_c_form"):
        col1, col2 = st.columns(2)
        with col1:
            app_name = st.text_input("द्वितीय अपीलकर्त्याचे नाव:", value="सतीश अशोक प्रधान")
            app_address = st.text_area("अपीलकर्त्याचा पत्ता:", value="छत्रपती संभाजीनगर")
            app_mobile = st.text_input("मोबाईल क्रमांक:", value="८६६८२३५३९५")
            bench = st.selectbox("राज्य माहिती आयोग खंडपीठ निवडा:", [
                "राज्य माहिती आयोग, राज्य खंडपीठ छत्रपती संभाजीनगर",
                "राज्य माहिती आयोग, मुख्य खंडपीठ मुंबई",
                "राज्य माहिती आयोग, खंडपीठ पुणे",
                "राज्य माहिती आयोग, खंडपीठ नागपूर",
                "राज्य माहिती आयोग, खंडपीठ नाशिक",
                "राज्य माहिती आयोग, खंडपीठ अमरावती"
            ])
        with col2:
            first_aa_details = st.text_area("प्रथम अपीलीय अधिकाऱ्याचा तपशील व निर्णय तारीख:", placeholder="मा. प्रथम अपीलीय अधिकारी, निर्णय तारीख...")
            pio_details = st.text_area("जन माहिती अधिकाऱ्याचा तपशील:", placeholder="जन माहिती अधिकारी कार्यालय...")
            
        penalty_demand = st.checkbox("माहिती अधिकार कायदा कलम २० अन्वये दोषी जन माहिती अधिकाऱ्यावर २५,००० रुपये दंडात्मक कारवाई व विभागीय चौकशीची मागणी करायची आहे का?", value=True)
        grounds = st.text_area("द्वितीय अपीलाची कायदेशीर कारणे व मुद्दे:", value="१. प्रथम अपीलीय अधिकाऱ्याने नैसर्गिक न्यायतत्त्वाचे पालन केले नाही.\n२. जन माहिती अधिकाऱ्याने जाणूनबुजून माहिती लपवली आहे.")

        submit_c = st.form_submit_button("🚀 जोडपत्र 'क' द्वितीय अपील तयार करा")

    if submit_c:
        penalty_text = ""
        if penalty_demand:
            penalty_text = "तसेच माहिती अधिकार अधिनियम २००५ च्या कलम २०(१) अन्वये दोषी जन माहिती अधिकाऱ्यावर रु. २५,०००/- (पंचवीस हजार) दंडात्मक कारवाई करण्यात यावी व कलम २०(२) अन्वये विभागीय शिस्तभंगाची कारवाईची शिफारस करण्यात यावी."

        rti_c_text = f"""परिशिष्ट / जोडपत्र 'क'
माहितीचा अधिकार अधिनियम, २००५ च्या कलम १९(३) अन्वये मा. राज्य माहिती आयोगाकडे द्वितीय अपील.

सामोर,
मा. राज्य माहिती आयुक्त,
{bench}

अपील अर्ज क्र. _________ / २०२५

{app_name},
वय: बाल/प्रौढ, व्यवसाय: नोंद नाही,
पत्ता: {app_address}, मो. {app_mobile}                  ... अपीलकर्ता

विरूद्ध

१. जन माहिती अधिकारी, {pio_details}                  ... प्रतिवादी क्र. १
२. प्रथम अपीलीय अधिकारी, {first_aa_details}            ... प्रतिवादी क्र. २

अपीलाचा विषय: कलम १९(३) अन्वये द्वितीय अपील अर्ज.

महोदय,
अपीलकर्ता खालीलप्रमाणे नमूद करतो की:

१. अपीलकर्त्याने जन माहिती अधिकाऱ्यांकडे कलम ६(१) अन्वये अर्ज सादर केला होता.
२. वेळेत माहिती न मिळाल्याने प्रथम अपीलीय अधिकाऱ्यांकडे कलम १९(१) अन्वये अपील दाखल केले.
३. परंतु, प्रतिवादींनी कायद्याचे उल्लंघन करून माहिती देण्याचे टाळले आहे.

अपीलाचे मुख्य आधार (Grounds):
{grounds}

मागणी / प्रार्थना (Prayer):
क) प्रतिवादी क्र. १ यास संपूर्ण माहिती विनामूल्य व प्रमाणित प्रतीसह देण्याचे आदेश व्हावेत.
ख) {penalty_text}
ग) अपीलकर्त्यास झालेल्या मानसिक व आर्थिक त्रासापोटी योग्य ती भरपाई देण्यात यावी.

सत्यप्रतिज्ञा
मी, {app_name}, शपथपूर्वक लिहून देतो की, वरील अपीलातील सर्व मजकूर माझ्या माहितीनुसार सत्य व बरोबर आहे.

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {app_name} )
                                                     अपीलकर्ता"""

        st.success("✅ जोडपत्र 'क' द्वितीय अपील तयार झाले आहे!")
        st.markdown(f'<div class="draft-container">{rti_c_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("जोडपत्र_क_द्वितीय_अपील", rti_c_text)

# ==============================================================================
# ९. विभाग ५: AI चॅट (Akanksha AI Assistant)
# ==============================================================================

elif st.session_state.current_page == "४) ✨ AI चॅट (आकांक्षा AI)":
    st.header("✨ आकांक्षा AI - कायदेशीर व प्रशासकीय चॅट सहाय्यक")
    st.info("💡 **टीप:** कोणत्याही कायदेशीर प्रश्नाचे, कलमांचे किंवा अर्जाचे स्पष्टीकरण मराठीत विचारा.")

    # चॅट हिस्ट्री डिस्प्ले
    for message in st.session_state.chat_history:
        if message["role"] == "user":
            st.markdown(f'<div class="user-bubble"><b>तुम्ही:</b><br>{message["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="ai-bubble"><b>🤖 आकांक्षा AI:</b><br>{message["content"]}</div>', unsafe_allow_html=True)

    st.write("")
    st.markdown("<div style='clear:both;'></div>", unsafe_allow_html=True)

    # इनपुट फॉर्म
    with st.form("chat_form", clear_on_submit=True):
        col_inp, col_btn = st.columns([8, 2])
        with col_inp:
            user_input = st.text_input("तुमचा प्रश्न किंवा कायदेशीर शंका येथे लिहा:", placeholder="उदा. रस्ता कामाच्या तक्रारीसाठी कोणता अर्ज करावा?")
        with col_btn:
            chat_submit = st.form_submit_button("संदेश पाठवा 🚀")

    if chat_submit and user_input:
        # युझर मेसेज सेव्ह
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        
        # AI रिस्पॉन्स सिम्युलेशन / कायदेशीर लॉजिक
        ai_response = f"माझ्या कायदेशीर विश्लेषानुसार: '{user_input}' या विषयासाठी तुम्ही आकांक्षा AI मधील शासकीय तक्रार अर्ज किंवा माहिती अधिकार (RTI) जोडपत्र 'अ' चा वापर करू शकता. यासंबंधी अधिक माहितीसाठी संबंधित विभागाच्या वरिष्ठ अधिकाऱ्यांकडे कलम/नियमानुसार तक्रार दाखल करावी."
        
        if "माहिती अधिकार" in user_input or "RTI" in user_input.upper():
            ai_response = "माहिती अधिकार कायदा २००५ च्या कलम ६(१) अन्वये तुम्ही कोणत्याही सरकारी कार्यालयाकडून माहिती मागवू शकता. अर्ज सादर केल्यापासून ३० दिवसांच्या आत माहिती मिळणे बंधनकारक आहे."
        elif "तक्रार" in user_input:
            ai_response = "प्रशासकीय गैरकारभाराविरुद्ध तक्रार करण्यासाठी तुम्ही जिल्हाधिकारी किंवा संबंधित विभागाच्या आयुक्तांकडे लेखी अर्ज करू शकता. आपल्या ॲपमधील 'शासकीय तक्रार अर्ज' हा पर्याय यासाठी सर्वोत्तम आहे."
            
        st.session_state.chat_history.append({"role": "assistant", "content": ai_response})
        st.rerun()

# ==============================================================================
# १०. विभाग ६: न्यायालयीन मसुदा (Court Drafts - Civil/Criminal/Writ)
# ==============================================================================

elif st.session_state.current_page == "५) 📜 न्यायालयीन मसुदा":
    st.header("📜 न्यायालयीन मसुदा - दिवाणी / फौजदारी / रिट याचिका")
    
    with st.form("court_draft_form"):
        col1, col2 = st.columns(2)
        with col1:
            court_name = st.text_input("न्यायालयाचे नाव:", value="मा. दिवाणी न्यायालय वरिष्ठ स्तर, छत्रपती संभाजीनगर")
            petitioner = st.text_input("वादी / याचिकाकर्त्याचे नाव:", value="सतीश अशोक प्रधान")
            petitioner_address = st.text_area("वादीचा पत्ता:", value="छत्रपती संभाजीनगर, मो. ८६६८२३५३९५")
            case_type = st.selectbox("दाव्याचा / याचिकेचा प्रकार:", ["दिवाणी दावा (Civil Suit)", "फौजदारी तक्रार (Criminal Complaint - Sec 156(3)/200)", "रिट याचिका (Writ Petition - Art 226/32)", "मनाई हुकूम अर्ज (Injunction Application)"])
        with col2:
            respondent = st.text_input("प्रतिवादीचे नाव:", placeholder="उदा. अ ब क / संबंधित विभाग")
            respondent_address = st.text_area("प्रतिवादीचा पत्ता:", placeholder="प्रतिवादीचा पूर्ण पत्ता...")
            cause_of_action = st.text_input("दाव्याचे कारण (Cause of Action) उद्भवल्याची तारीख:", placeholder="उदा. १० जानेवारी २०२५ रोजी")

        facts = st.text_area("दाव्याची / प्रकरणाची सविस्तर वस्तुस्थिती (Facts):", placeholder="१. वादी हा सदर मालमत्तेचा कायदेशीर मालक आहे...\n२. प्रतिवादीने बेकायदेशीरपणे हस्तक्षेप करण्याचा प्रयत्न केला...")
        legal_points = st.text_area("कायदेशीर मुद्दे व कलमे (Legal Grounds):", placeholder="उदा. भारतीय न्याय संहिता (BNS) / Specific Relief Act च्या कलमांनुसार...")
        prayer = st.text_area("मा. न्यायालयाकडे केलेली अंतिम मागणी (Prayer):", placeholder="अ) वादीच्या बाजूने आणि प्रतिवादीच्या विरोधात कायमस्वरूपी मनाई हुकूमाचा हुकूमनामा व्हावा.")

        submit_court = st.form_submit_button("🚀 न्यायालयीन मसुदा तयार करा")

    if submit_court:
        court_text = f"""इन द कोर्ट ऑफ {court_name}

दावा / याचिका प्रकार: {case_type}
दावा क्र. ________ / २०२५

{petitioner}
वय: ____, व्यवसाय: __________,
पत्ता: {petitioner_address}                        ... वादी / याचिकाकर्ता

विरूद्ध

{respondent}
वय: ____, व्यवसाय: __________,
पत्ता: {respondent_address}                        ... प्रतिवादी

विषय: {case_type} - अंतरीम व अंतिम दिलासा मिळणेबाबत.

मा. न्यायालयात वादीचा निवेदन अर्ज खालीलप्रमाणे आहे:

१. वादी {petitioner_address} येथील रहिवासी असून सदर प्रकरणातील बाधित पक्षकार आहे.

२. प्रकरणाची वस्तुस्थिती (Facts of the Case):
{facts}

३. दाव्याचे कारण (Cause of Action):
सदर दाव्याचे कारण दिनांक {cause_of_action} रोजी उद्भवले असून सदर न्यायालयीन हद्दीतच घटना घडल्यामुळे मा. न्यायालयास या दाव्याची सुनावणी करण्याचा पूर्ण अधिकार आहे.

४. कायदेशीर आधार व मुद्दे (Legal Grounds):
{legal_points}

५. प्रार्थना / मागणी (Prayer):
वादीची मा. न्यायालयाकडे अशी विनंती आहे की:
{prayer}
ब) या दाव्याचा होणारा संपूर्ण खर्च प्रतिवादीकडून वादीस देववण्यात यावा.
क) मा. न्यायालयास या प्रकरणाच्या न्याय्य हक्कासाठी जे योग्य वाटेल ते अन्य आदेश व्हावेत.

सत्यप्रतिज्ञा (Verification)
मी, {petitioner}, प्रतिज्ञापूर्वक नमूद करतो की, वरील परिच्छेद १ ते ५ मधील मजकूर माझ्या स्वतःच्या माहितीनुसार आणि समजुतीनुसार सत्य व बरोबर आहे.

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {petitioner} )
                                                     वादी / याचिकाकर्ता"""

        st.success("✅ न्यायालयीन मसुदा तयार झाला आहे!")
        st.markdown(f'<div class="draft-container">{court_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("न्यायालयीन_मसुदा", court_text)

# ==============================================================================
# ११. विभाग ७: शासकीय तक्रार अर्ज (Govt Grievance Application)
# ==============================================================================

elif st.session_state.current_page == "६) शासकीय तक्रार अर्ज":
    st.header("📢 शासकीय तक्रार अर्ज - प्रशासकीय गैरकारभाराविरुद्ध")
    
    with st.form("grievance_form"):
        col1, col2 = st.columns(2)
        with col1:
            app_name = st.text_input("तक्रारदाराचे नाव:", value="सतीश अशोक प्रधान")
            app_address = st.text_area("तक्रारदाराचा पत्ता:", value="छत्रपती संभाजीनगर, मो. ८६६८२३५३९५")
            target_officer = st.text_input("ज्यांच्याकडे तक्रार करायची आहे ते अधिकारी:", value="मा. जिल्हाधिकारी साहेब / मा. आयुक्त साहेब")
        with col2:
            dept_name = st.text_input("संबंधित विभाग / कार्यालय:", placeholder="उदा. सार्वजनिक बांधकाम विभाग / महानगरपालिका")
            against_whom = st.text_input("ज्या अधिकारी/कर्मचाऱ्याविरुद्ध तक्रार आहे त्याचे नाव/पद:", placeholder="उदा. संबंधित कनिष्ठ अभियंता / अधिकारी")

        subject = st.text_input("तक्रारीचा विषय:", placeholder="उदा. रस्त्याच्या निकृष्ट दर्जाच्या कामाबाबत व भ्रष्टाचाराची चौकशी करणेबाबत.")
        grievance_details = st.text_area("तक्रारीचे सविस्तर वर्णन व गैरकारभाराचा तपशील:", placeholder="१. सदर कामात निकृष्ट दर्जाचे साहित्य वापरले गेले आहे.\n२. नागरिकांच्या सुरक्षेशी खेळ केला जात आहे.")
        action_demanded = st.text_area("अपेक्षित प्रशासकीय कारवाई:", value="सदर प्रकरणाची निष्पक्ष चौकशी करून दोषींवर तात्काळ निलंबित करण्याची कारवाई करण्यात यावी व कामाचा दर्जा सुधारण्यात यावा.")

        submit_grievance = st.form_submit_button("🚀 शासकीय तक्रार अर्ज तयार करा")

    if submit_grievance:
        grievance_text = f"""प्रति,
{target_officer},
{dept_name},
छत्रपती संभाजीनगर.

विषय: {subject}

तक्रारदार : {app_name}
पत्ता      : {app_address}

महोदय,

मी खालील स्वाक्षरीकरणार आपणास विनंतीपूर्वक अर्ज सादर करतो की:

१. मी वरील नमूद पत्यावरील रहिवासी असून सदर भागातील प्रशासकीय गैरकारभार व समस्येकडे आपले लक्ष वेधू इच्छितो.

२. तक्रारीचा सविस्तर तपशील:
सदर {dept_name} अंतर्गत {against_whom} यांच्या दुर्लक्षामुळे व गैरकारभारामुळे नागरिकांना मोठ्या त्रासाला सामोरे जावे लागत आहे.
तपशील खालीलप्रमाणे:
{grievance_details}

३. माझी आपणास नम्र विनंती आहे की, या अर्जाची गांभीर्याने दखल घेऊन:
{action_demanded}

४. जर या तक्रारीवर पुढील १५ दिवसांत योग्य ती कायदेशीर व प्रशासकीय कारवाई न झाल्यास, मला सनदशीर मार्गाने तीव्र आंदोलन व वरिष्ठ पातळीवर दाद मागावी लागेल याची नोंद घ्यावी.

आपला विश्वासू,

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {app_name} )
                                                     तक्रारदार"""

        st.success("✅ शासकीय तक्रार अर्ज तयार झाला आहे!")
        st.markdown(f'<div class="draft-container">{grievance_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("शासकीय_तक्रार_अर्ज", grievance_text)

# ==============================================================================
# १२. विभाग ८: प्रतिज्ञापत्र (Affidavit Draft)
# ==============================================================================

elif st.session_state.current_page == "७) 📝 प्रतिज्ञापत्र (Affidavit)":
    st.header("📝 प्रतिज्ञापत्र (Affidavit Draft)")
    
    with st.form("affidavit_form"):
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("प्रतिज्ञापत्र देणाऱ्याचे पूर्ण नाव:", value="सतीश अशोक प्रधान")
            age = st.text_input("वय:", placeholder="उदा. ३५")
            occupation = st.text_input("व्यवसाय:", placeholder="उदा. व्यवसाय / नोकरी")
            address = st.text_area("पूर्ण पत्ता:", value="छत्रपती संभाजीनगर")
        with col2:
            purpose = st.selectbox("प्रतिज्ञापत्राचे कारण / उद्देश:", [
                "नाव / नावात बदल असलेबाबत प्रतिज्ञापत्र",
                "पत्याचा पुरावा / रहिवासी प्रतिज्ञापत्र",
                "आर्थिक उत्पन्नाबाबत प्रतिज्ञापत्र",
                "शासकीय योजनेच्या लाभासाठी प्रतिज्ञापत्र",
                "कागदपत्रे गहाळ झालेबाबत प्रतिज्ञापत्र",
                "इतर सामान्य प्रतिज्ञापत्र (General Affidavit)"
            ])
            id_proof = st.text_input("ओळख पुरावा (आधार क्र. / पॅन क्र.):", placeholder="उदा. आधार क्र. १२३४-५६७८-९०१२")

        statements = st.text_area("प्रतिज्ञापत्रातील मुख्य सत्य विधाने (प्रत्येक विधान नवीन ओळीवर):", placeholder="१. मी असे प्रतिज्ञापूर्वक सांगतो की माझे मूळ नाव सतीश अशोक प्रधान आहे.\n२. माझ्या सर्व सरकारी कागदपत्रांवर हेच नाव अचूक आहे.")

        submit_affidavit = st.form_submit_button("🚀 प्रतिज्ञापत्र मसुदा तयार करा")

    if submit_affidavit:
        affidavit_text = f"""प्रतिज्ञापत्र (AFFIDAVIT)

मी, {name}, वय: {age} वर्षे, व्यवसाय: {occupation},
रा. {address}, ओळख पुरावा आधार क्र./पॅन क्र.: {id_proof},
खालीलप्रमाणे प्रतिज्ञापूर्वक लिहून देतो की:

१. मी वरील नमूद पत्यावर कायमस्वरूपी राहणारा असून भारताचा नागरिक आहे.

२. सदर प्रतिज्ञापत्र मी '{purpose}' या कारणासाठी करत आहे.

३. मी शपथपूर्वक खालीलप्रमाणे सत्य विधाने घोषित करतो:
{statements}

४. वरील दिलेली सर्व माहिती व विधाने माझ्या वैयक्तिक माहितीनुसार व विश्वासानुसार पूर्णतः सत्य व बरोबर आहेत. त्यात कोणताही फरक किंवा असत्य माहिती दडवलेली नाही.

५. जर भविष्यात वरील माहिती खोटी किंवा चुकीची आढळल्यास, त्यास मी वैयक्तिकरीत्या भारतीय न्याय संहिता (BNS) व संबंधित कायद्यानुसार शिक्षेस पात्र राहीन.

करिता हे प्रतिज्ञापत्र आज दिनांक {datetime.datetime.now().strftime('%d/%m/%Y')} रोजी छत्रपती संभाजीनगर येथे लिहून दिले.

प्रतिज्ञापत्र देणारा:

                                                    ( {name} )
                                                     स्वाक्षरी / अंगठा

सत्यता पडताळणी (Verification)
माझ्यासमोर {name} यांनी उपस्थित राहून मजकूर वाचून दाखवल्यानंतर सत्य असल्याचे मान्य करून स्वाक्षरी केली.

दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}
ठिकाण: छत्रपती संभाजीनगर                      मा. कार्यकारी दंडाधिकारी / नोटीरी"""

        st.success("✅ प्रतिज्ञापत्र मसुदा तयार झाला आहे!")
        st.markdown(f'<div class="draft-container">{affidavit_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("प्रतिज्ञापत्र_Affidavit", affidavit_text)

# ==============================================================================
# १३. विभाग ९: ग्राहक मंच तक्रार (Consumer Forum Complaint)
# ==============================================================================

elif st.session_state.current_page == "८) 🛒 ग्राहक मंच तक्रार":
    st.header("🛒 ग्राहक मंच तक्रार (Consumer Protection Act 2019)")
    
    with st.form("consumer_form"):
        col1, col2 = st.columns(2)
        with col1:
            complainant = st.text_input("ग्राहकाचे (तक्रारदाराचे) नाव:", value="सतीश अशोक प्रधान")
            complainant_address = st.text_area("ग्राहकाचा पत्ता:", value="छत्रपती संभाजीनगर, मो. ८६६८२३५३९५")
            forum_name = st.text_input("ग्राहक मंचाचे नाव:", value="मा. जिल्हा ग्राहक विवाद निवारण आयोग, छत्रपती संभाजीनगर")
        with col2:
            opposite_party = st.text_input("विरोधी कंपनी / विक्रेत्याचे नाव:", placeholder="उदा. मे. एक्स वाय झेड इलेक्ट्रॉनिक्स / कंपनी")
            opposite_address = st.text_area("विरोधी पक्षाचा पत्ता:", placeholder="कंपनीचा पूर्ण पत्ता...")
            product_service = st.text_input("खरेदी केलेली वस्तू / सेवा:", placeholder="उदा. मोबाईल / वॉशिंग मशीन / इन्शुरन्स पोलिसी")

        purchase_date = st.text_input("खरेदीची तारीख व बिल रक्कम:", placeholder="उदा. १५ ऑगस्ट २०२४, रक्कम रु. ४५,०००/-")
        defects = st.text_area("सेवेतील त्रुटी / वस्तूतील दोष (Defect/Deficiency):", placeholder="१. वस्तू वारंवार खराब होत आहे.\n२. वॉरंटी असतानाही सर्व्हिस देण्यास नकार दिला.")
        compensation = st.text_area("मागणी व भरपाईची रक्कम (Compensation Claimed):", placeholder="१. मूळ रक्कम रु. ४५,०००/- परतावा मिळावा.\n२. मानसिक त्रासापोटी रु. २५,०००/- भरपाई मिळावी.")

        submit_consumer = st.form_submit_button("🚀 ग्राहक मंच तक्रार तयार करा")

    if submit_consumer:
        consumer_text = f"""सामोर,
{forum_name}

तक्रार अर्ज क्र. ________ / २०२५
(ग्राहक संरक्षण कायदा, २०१९ च्या कलम ३५ अन्वये)

{complainant},
रा. {complainant_address}                            ... तक्रारदार / ग्राहक

विरूद्ध

{opposite_party},
पत्ता: {opposite_address}                            ... विरोधी पक्ष / विक्रेता

विषय: सेवेतील त्रुटी (Deficiency of Service) व सदोष वस्तूंबाबत तक्रार अर्ज.

तक्रारदाराचा अर्ज खालीलप्रमाणे आहे:

१. तक्रारदार हा ग्राहक संरक्षण कायदा २०१९ च्या कलम २(७) नुसार विरोधी पक्षाचा ग्राहक आहे.

२. तक्रारदाराने विरोधी पक्षाकडून {product_service} ची खरेदी {purchase_date} रोजी केली होती.

३. सेवेतील त्रुटी व दोषांचा तपशील:
{defects}

४. विरोधी पक्षाने ग्राहकास दिलेली सेवा अत्यंत दर्जाहीन असून ग्राहकाची फसवणूक केली आहे. त्यामुळे ग्राहक संरक्षण कायद्याचे उल्लंघन झाले आहे.

५. मागणी व प्रार्थना (Relief Claimed):
तक्रारदार मा. आयोगाकडे अशी प्रार्थना करतो की:
{compensation}
३) या तक्रारीचा अर्ज खर्च रु. ५,०००/- विरोधी पक्षाकडून देववण्यात यावा.

सत्यप्रतिज्ञा
मी, {complainant}, प्रतिज्ञापूर्वक सांगतो की वरील सर्व मजकूर माझ्या माहितीनुसार सत्य आहे.

ठिकाण: छत्रपती संभाजीनगर
दिनांक: {datetime.datetime.now().strftime('%d/%m/%Y')}

                                                    ( {complainant} )
                                                     तक्रारदार"""

        st.success("✅ ग्राहक मंच तक्रार मसुदा तयार झाला आहे!")
        st.markdown(f'<div class="draft-container">{consumer_text}</div>', unsafe_allow_html=True)
        render_draft_action_buttons("ग्राहक_मंच_तक्रार", consumer_text)

# ==============================================================================
# १४. विभाग १०: कायदेशीर संदर्भ ग्रंथालय (Legal Reference Library)
# ==============================================================================

elif st.session_state.current_page == "📚 कायदेशीर संदर्भ ग्रंथालय":
    st.header("📚 कायदेशीर संदर्भ ग्रंथालय (Legal Reference Library)")
    st.write("येथे महत्त्वाचे भारतीय कायदे, कलमे व नागरिकांच्या अधिकारांची माहिती दिली आहे:")

    tab1, tab2, tab3, tab4 = st.tabs(["📄 माहिती अधिकार (RTI)", "🛒 ग्राहक संरक्षण", "⚖️ BNS / IPC कलमे", "🏛️ प्रशासकीय नियम"])

    with tab1:
        st.subheader("माहिती अधिकार अधिनियम, २००५ - महत्त्वाची कलमे")
        st.markdown("""
        * **कलम ६(१):** माहिती मिळवण्यासाठी जन माहिती अधिकाऱ्याकडे मूळ अर्ज करणे.
        * **कलम ६(३):** अर्ज चुकीच्या खात्याकडे गेल्यास ५ दिवसांच्या आत योग्य विभागाकडे वर्ग करणे.
        * **कलम ७(१):** ३० दिवसांच्या आत माहिती देणे बंधनकारक (जीव व स्वातंत्र्याशी संबंधित असल्यास ४८ तासांत).
        * **कलम ७(६):** ३० दिवसांत माहिती न दिल्यास ती **विनामूल्य (Free of Cost)** देणे बंधनकारक.
        * **कलम ८:** माहिती देण्यापासून सूट (देशाची सुरक्षा, गोपनीय बाबी).
        * **कलम १९(१):** ३० दिवसांत माहिती न मिळाल्यास प्रथम अपीलीय अधिकाऱ्याकडे प्रथम अपील.
        * **कलम १९(३):** राज्य माहिती आयोगाकडे ९० दिवसांत द्वितीय अपील.
        * **कलम २०(१):** माहिती नाकारणाऱ्या अधिकाऱ्यावर **दररोज रु. २५० ते कमाल रु. २५,००० दंडाची तरतूद**.
        """)

    with tab2:
        st.subheader("ग्राहक संरक्षण कायदा, २०१९ - महत्त्वाची माहिती")
        st.markdown("""
        * **जिल्हा आयोग (District Commission):** रु. ५० लाखांपर्यंतच्या दाव्यांसाठी.
        * **राज्य आयोग (State Commission):** रु. ५० लाख ते रु. २ कोटींपर्यंतच्या दाव्यांसाठी.
        * **राष्ट्रीय आयोग (National Commission):** रु. २ कोटींपेक्षा जास्त दाव्यांसाठी.
        * **तक्रार मुदत:** कारण उद्भवल्यापासून **२ वर्षांच्या आत** तक्रार दाखल करणे आवश्यक.
        """)

    with tab3:
        st.subheader("भारतीय न्याय संहिता (BNS) / IPC मुख्य कलमे")
        st.markdown("""
        * **फसवणूक (Cheating):** BNS Sec 318 (Old IPC 420).
        * **बनावट कागदपत्रे (Forgery):** BNS Sec 336 (Old IPC 465).
        * **बेकायदेशीर अडवणूक:** BNS Sec 126 (Old IPC 341).
        * **मानहानी (Defamation):** BNS Sec 356 (Old IPC 499/500).
        """)

    with tab4:
        st.subheader("प्रशासकीय तक्रार व नागरिक सनद")
        st.markdown("""
        * **नागरी सेवा हक्क कायदा:** शासकीय सेवा ठराविक मुदतीत मिळण्याचा कायदेशीर अधिकार.
        * **दप्तर दिरंगाई कायदा:** शासकीय नस्ती (Files) ७ दिवसांपेक्षा जास्त काळ प्रलंबित ठेवल्यास दंडात्मक कारवाई.
        """)

# ==============================================================================
# १५. विभाग ११: जतन केलेले मसुदे (Draft History Manager)
# ==============================================================================

elif st.session_state.current_page == "💾 माझे जतन केलेले मसुदे":
    st.header("💾 माझे जतन केलेले मसुदे (Draft History)")

    if not st.session_state.draft_history:
        st.info("ℹ️ सध्या कोणताही मसुदा सेव्ह केलेला नाही. नवीन मसुदा तयार केल्यानंतर 'मसुदा इतिहासात जतन करा' बटणावर क्लिक करा.")
    else:
        st.write(f"एकूण जतन केलेले मसुदे: **{len(st.session_state.draft_history)}**")
        
        for idx, item in enumerate(reversed(st.session_state.draft_history)):
            with st.expander(f"📌 [{item['date']}] {item['type']} (आयडी: #{item['id']})"):
                st.markdown(f'<div class="draft-container">{item["content"]}</div>', unsafe_allow_html=True)
                
                col_h1, col_h2 = st.columns(2)
                with col_h1:
                    docx_f = generate_docx_file(item["content"], title=item["type"])
                    st.download_button(
                        label="📄 Word डाऊनलोड",
                        data=docx_f,
                        file_name=f"{item['type']}_{item['id']}.docx",
                        key=f"hist_docx_{idx}"
                    )
                with col_h2:
                    pdf_f = generate_pdf_file(item["content"], title=item["type"])
                    st.download_button(
                        label="🔴 PDF डाऊनलोड",
                        data=pdf_f,
                        file_name=f"{item['type']}_{item['id']}.pdf",
                        key=f"hist_pdf_{idx}"
                    )

        st.write("")
        if st.button("🗑️ सर्व इतिहास पुसून टाका (Clear History)"):
            st.session_state.draft_history = []
            st.success("सर्व मसुदा इतिहास क्लिअर झाला आहे.")
            st.rerun()

# ==============================================================================
# १६. फुटर (Footer)
# ==============================================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #64748B; font-size: 0.9rem; padding: 10px;">
    <b>आकांक्षा AI कायदेशीर व प्रशासकीय महा-सहाय्यक</b> © २०२६ | संकल्पना व निर्मिती: <b>सतीश अशोक प्रधान</b> (मो. ८६६८२३५३९५)<br>
    <i>सर्व हक्क सुरक्षित. हे ॲप्लिकेशन कायदेशीर माहिती व मसुदा सहाय्यासाठी बनवले आहे.</i>
</div>
""", unsafe_allow_html=True)
